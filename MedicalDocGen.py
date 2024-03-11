# you do NOT have to memorise "how to perfectly write code from your brain" just understand the concepts
# most of the time, you will need to google the structure and then adapt it to your code 
# it's like if google is your Maths formula cheat sheet instead of memorising all of the maths formulas

# EXTRA INFO: 

  # A "VARIABLE" is a "Temporary storage of a value"
  # if you see anything like [ keyWord = "value" ] then this is when we are temporarily setting a value
  # you will want to do this when you need to make your code "dynamic" and change depending on the situation
  # you will also want to do this if you need to keep referring the the "variable" for a calculation for example
  # variables can store either A SINGLE VALUE or MANY VALUES.  
    # Single value example: x = 1
    # Multiple value example x = [1,2,3,4,5] <- this is called an array
  # in this code, you'll see that I use "Dictionaries" a lot. See below for more info

  # A "DICTIONARY" and is structured like:
  # [ dictionary = {"SUBJECT": "VALUE", "SUBJECT2": "VALUE2"} ]
  # INFO: THE COMMA IS USED IF YOU WANT MULTIPLE SUBJECTS & VALUES 
  # TO REFER TO A SPECIFIC SUBJECT IN YOUR DICTIONARY, THEN YOU FORMAT IT LIKE:
  # "dictionary['Subject']" AND THEN YOU WILL RECEIVE THE CORRELATED "VALUE"

  # A "FUNCTIONS" is like "mini codes" that you can re-use in your "General code"
  # Why use functions? If you want to keep your code clean and keep re-using a specific "mini code" then you want to use a function
  # in python, functions begin like below:

    # def doSomething(parameter):
    #    print("this function is doing something")
    #    print(f"This is a parameter: {parameter})

  # IF you need to "call" the function or trigger the function, then somewhere else in your code you just need to refer to it like:
    # doSomething([INSERT PARAMETER VALUE HERE])


# INFO: by default, python only has a "skeleton" amount of functions that it can run 
# INFO: you need to "import libraries" so that your python code can do more things
# INFO: BELOW we import a few libraries so that our python code can:
  # Create a word doc: docx
  # do calculations on date/time: datetime, pytz,calendar

import docx
# LIBRARY TO GET DATE/TIME
from datetime import datetime, timedelta, timezone, date
# LIBRARY FOR TIMEZONE
import pytz
# LIBRARY TO CONVERT NUMBER TO WEEKDAY NAME
import calendar

# DEFINE THE CURRENT DATE & TIME - I just stole this code from the internet to get Current Sydney time
current_date = date.today().strftime("%Y-%m-%d")
sydney_tz = pytz.timezone('Australia/Sydney')
current_datetime = datetime.now(sydney_tz).strftime("%Y-%m-%d-%H-%M-%S")
sydney_date = datetime.now(pytz.timezone('Australia/Sydney'))



# PRE-DEFINING A LIST OF QUESTIONS - ASKING THE USER QUESTIONS TO GET INFORMATION 
def getQuestions():

  # BELOW IS CALLED A "DICTIONARY" - SEE INFORMATION ABOUT "DICTIONARY" AT THE TOP OF THE CODE
  questions = { 
      "Enter the heart rate": 0
      , "Enter the client's Name (First & Last Name)": ""
      , "Enter the client's appointment time (24 HOUR TIME HHMM)": ""
      , "Enter the client's appointment date (YYYY-MM-DD)": ""
  }
  


  # LOOP THROUGH THE DICTIONARY LIST AND ADD EACH ANSWER INTO AN ARRAY
  for i in questions: # BEGIN THE LOOP AND GO THROUGH ALL OF THE QUESTIONS IN THE DICTIONARY
    question = i # SET THE "question" TO BE THE CURRENT LOOP VALUE
    temp_var = input(f"{question}: ") # THIS "input" KEYWORD ALLOWS PYTHON TO TAKE AN INPUT FROM THE USER
    questions[i] = temp_var # SETTING THE "VALUE" OF THE DICTIONARY SUBJECT TO EQUAL WHAT THE USER ENTERED
    # E.G. if the question is "What is the heart rate?" the above line will set the value to be the user input

  # DEPENDING ON THE HEART RATE, THEN PROVIDE THE USER WITH A DIFFERENT TEMPLATE - 90 IS HARD CODED 
  if int(questions['Enter the heart rate']) >= 90: #CHECKING IF THE USER INPUTTED 90 OR GREATER
      template_version = "A" # SET THE TEMPLATE TO BE "A" 
  elif int(questions['Enter the heart rate']) < 90: #ELSE IF THE HEART RATE IS LOWER THAN 90, THEN DO SOMETHING ELSE
      template_version = "B" # SET THE TEMPLATE TO BE "B" 

  # SEE EXTRA INFO AT THE TOP ABOUT "FUNCTIONS"
      
  # CALL FUNCTIONS TO FORMAT DATE VALUES
  # IN THE BELOW PARAMETER, WE PASS THROUGH THE DATE ENTERED BY THE USER
  # WE ALSO PASS IN THE TYPE OF DATA WE WANT TO TAKE BACK.
  # to follow the flow of the data, look for "def formatDate"
  # at the end of the "def formatDate" you will need to come back here to continue the code sequence
  week_day_num = formatDate(questions["Enter the client's appointment date (YYYY-MM-DD)"], "week_day_num")
  week_day = formatDate(questions["Enter the client's appointment date (YYYY-MM-DD)"], "week_day")
    
  # CALL FUNCTION TO BEGIN CREATING THE DOCUMENT
  # we also "PASS IN" 4 varaibles, questions, template_version, week_day, week_day_num
  createDocument(questions, template_version, week_day, week_day_num)

# BEGINING OF FUNCTION TO FORMAT THE DATE VALUES
def formatDate(date_string, requested_value): 
  # INITIALISING VARIABLE
  # Initialising also means "defining" or "creating" the variable so that your code knows it should exist
  return_value = ""
  
  # DATE CONVERSION INTO PROPER FORMAT
  date_formatted = datetime.strptime(date_string, '%Y-%m-%d') #strptime is a python defined function. we only understand how to use it by reading documentation online
  
  # USING THE FORMATTED DATE TO GET WEEK DAY NUMBER
  week_day_num = date_formatted.weekday()

  # USING THE WEEKDAY NUMBER TO GET THE WEEK DAY NAME
  week_day = calendar.day_name[date_formatted.weekday()]
  
  # DEPENDING ON WHAT VARIABLE/PARAMETER IS PASSED IN, THEN RETURN A DIFFERENT VALUE (NUMBER OR NAME)
  if requested_value == "week_day":
    return_value = week_day
    print(f"Formatted Date: {date_formatted}, Weekday: {week_day}")
  elif requested_value == "week_day_num":
    return_value = week_day_num
    print(f"Formatted Date: {date_formatted},  Week day num: {week_day_num}")
  
  # WE RETURN "return_value" BACK TO THE week_day_num & week_day ON LINE 79 & 80
  return return_value



        
# CREATE THE DOCUMENT NOTICE THAT THERE ARE 4 PARAMETERS
def createDocument(questions, template_version, week_day, week_day_num):
  # Calling the "calculateDates" function by passing in week_day_num
  # look for "def calculateDates"
  minus_days = calculateDates(week_day_num)
  # Calling the "calculateTime" function by passing in the user inputted time
  # look for "def calculateTime"
  calculated_time = calculateTime(questions["Enter the client's appointment time (24 HOUR TIME HHMM)"])
  
  #I STOLE MOST OF THE BELOW CODE FROM THE INTERNET, JUST TRUST THAT THEY DO WHAT YOU NEED
  # Create a document
  doc = docx.Document()

  # Add a paragraph to the document
  p = doc.add_paragraph()

  # Add some formatting to the paragraph
  p.paragraph_format.line_spacing = 1
  p.paragraph_format.space_after = 0

  # Add a run to the paragraph
  # run = p.add_run("python-docx")
  run = p.add_run("")

  # Add some formatting to the run
  run.bold = True
  run.italic = True
  run.font.name = 'Arial'
  run.underline
  run.font.size = docx.shared.Pt(16)

  # Add more text to the same paragraph
  run = p.add_run("CARDIAC CT PREPARATION")

  # paragraph.alignment = 0 # for left, 1 for center, 2 right, 3
  # Format the run
  run.alignment = 1
  run.bold = True
  run.font.name = 'Arial'
  run.font.size = docx.shared.Pt(16)

  # Add another paragraph (left blank for an empty line)
  doc.add_paragraph()

  # Add another paragraph
  p = doc.add_paragraph()
  run.font.size = docx.shared.Pt(12)

  # Add a run and format it
  # USING THE USER INPUTS TO FILL IN WORDS FOR THE WORD DOCUMENT
  var_name = questions["Enter the client's Name (First & Last Name)"]
  run = p.add_run(f"Name: {var_name}")
  run.font.name = 'Arial'
  run.font.size = docx.shared.Pt(12)

  p = doc.add_paragraph()
  # USING THE USER INPUTS TO FILL IN WORDS FOR THE WORD DOCUMENT
  var_date = questions["Enter the client's appointment date (YYYY-MM-DD)"]
  var_time = questions["Enter the client's appointment time (24 HOUR TIME HHMM)"]
  run = p.add_run(f"Appointment: {week_day} {var_date} at {calculated_time['appointment_time']}. Arrive 15 minutes before at {calculated_time['arrival_time']}")
  run.font.name = 'Arial'
  run.font.size = docx.shared.Pt(12)

  doc.add_paragraph()

  # BASED ON THE USER'S HEART RATE... A OR B, THEN APPLY A DIFFERENT TEMPLATE

  if template_version == "A": # THIS IS OPTION A
    p = doc.add_paragraph()
    run = p.add_run(f"""
Take one tablet (Metoprolol 50mg) {minus_days['1_days_before']} night before bed.

Take one tablet (Metoprolol 50mg) {week_day} Morning at {calculated_time['2_hours_before_appointment_time']}, 2 hours before appointment time.
    """)
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
  elif template_version == "B": #THIS IS OPTION B
    p = doc.add_paragraph()
    run = p.add_run(f"YOU'RE HEALTHY, DON'T NEED TO TAKE ANY MEDICINE.")
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
  # THE IF ELSE STATEMENT ENDS HERE
    
  #CONTINUE ADDING INFORMATION TO THE BOTTOM OF THE DOCUMENT
  doc.add_paragraph()
  p = doc.add_paragraph()
  run = p.add_run(f"{minus_days['1_days_before']}")
  run.font.name = 'Arial'
  run.underline = True
  run.font.size = docx.shared.Pt(12)

  p = doc.add_paragraph()
  run = p.add_run(f"""
    No caffeine or stimulants (Coffee, tea, diet pills, sports drinks, chocolate)
    No Viagra""")
  run.font.name = 'Arial'
  run.font.size = docx.shared.Pt(12)

  doc.add_paragraph()
  p = doc.add_paragraph()
  run = p.add_run(f"{week_day} Morning")
  run.font.name = 'Arial'
  run.underline = True
  run.font.size = docx.shared.Pt(12)

  p = doc.add_paragraph()
  run = p.add_run("""
    No exercise
    No smoking or anti smoking medicines
    Fasting for four hours prior to your appointment
    Keep up water intake
    """)
  run.font.name = 'Arial'
  run.underline
  run.font.size = docx.shared.Pt(12)

  p = doc.add_paragraph()
  run = p.add_run("Take all your normal medications")
  run.font.name = 'Arial'
  run.underline
  run.font.size = docx.shared.Pt(12)

  # Save the document
  client_name = questions["Enter the client's Name (First & Last Name)"]
  doc_name = (f"South-East-Radiology-Instructions-{client_name}-{current_datetime}.docx").replace(" ","-")
  doc.save(doc_name)

# START OF DATE CALCULATION FUNCTION
# DEPENDING ON THE CALCULATED VALUE APPLY A TRANSFORMATION TO GET THE CORRECT DATE
def calculateDates(week_day_num):
  #CREATING ANOTHER DICTIONARY TO STORE THE CALCULATED DAYS
  returnDayNums ={
      "2_days_before": ""
      ,"1_days_before": ""
  }

  # DOING SOME MATHS TO UNDERSTAND THE DAY BEFORE THE APPOINTMENT DATE
  returnDayNums["2_days_before"] = week_day_num -2
  returnDayNums["1_days_before"] = week_day_num -1

  # IF THE VALUE FALLS UNDER 0, THEN YOU NEED TO ADD 7 TO BRING IT BACK INTO THE DATE RANGE
  if returnDayNums["2_days_before"] < 0:
      returnDayNums["2_days_before"] = calendar.day_name[returnDayNums["2_days_before"]+7]
  else: 
      returnDayNums["2_days_before"] =calendar.day_name[returnDayNums["2_days_before"]]

  if returnDayNums["1_days_before"] < 0:
      returnDayNums["1_days_before"] = calendar.day_name[returnDayNums["1_days_before"]+7]
  else:
      returnDayNums["1_days_before"] = calendar.day_name[returnDayNums["1_days_before"]]
  
  print(returnDayNums)
  return returnDayNums # RETURNING THE VALUE BACK TO LINE 121

# CALCULATING THE APPOINTMENT TIME AND ARRIVAL TIME
def calculateTime(time):
  #CREATING A DICTIONARY THAT WILL HOLD 3 VALUES AS EXPLAINED BY THE NAME
  formatted_time = {
    "appointment_time": ""
    ,"arrival_time": ""
    , "2_hours_before_appointment_time": ""
  }
  
  # CHANGING THE USER'S INPUT OF "HHMM" TO A PROPER "TIME OBJECT" THAT PYTHON CAN DO CALCULATIONS ON
  converted_appointment_time = datetime.strptime(time, "%H%M")

  # CHANGING THE USER'S INPUT OF "HHMM" TO A PYTHON TIME OBJECT AND THEN TAKING 15 MINUTES EARLIER
  converted_arrival_time = datetime.strptime(time, "%H%M") - timedelta(minutes = 15)

  # CHANGING THE USER'S INPUT OF "HHMM" TO A PYTHON TIME OBJECT AND THEN TAKING 2 HOURS
  converted_2_hours_before_appointment_time = datetime.strptime(time, "%H%M") - timedelta(hours = 2)

  # SETTING THE VALUES FROM THE DICTIONARY TO EQUAL THE NEW CALCULATED TIMES
  formatted_time["appointment_time"] = converted_appointment_time.strftime("%I:%M %p")
  formatted_time["arrival_time"] = converted_arrival_time.strftime("%I:%M %p")
  formatted_time["2_hours_before_appointment_time"] = converted_2_hours_before_appointment_time.strftime("%I:%M %p")
  print(formatted_time)

  # RETURNING THE DICTIONARY BACK TO LINE 124 SO WE CAN USE IT IN THE WORD DOCUMENT
  return formatted_time

# TYPICAL PYTHON CODE THAT WILL TELL IT WHERE THE BEGIN THE CODE SEQUENCING.
def main():
  getQuestions() # START EVERYTHING FROM: "def getQuestions():?""

if __name__ == '__main__':
    main()